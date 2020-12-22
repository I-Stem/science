import docx
import os
import base64
import pandas as pd
from io import StringIO, BytesIO
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING
from typing import Dict
import enchant
import string as string_library 
from .html_parser import html_formating
import time

def list_formatting ( line_list, level=1 ) :

    text = ""
    bullet_string = ''
    idx = 0

    while ( idx < len(line_list) ) :
        line = line_list[idx]
        line_text =  " ".join( [st["content"].strip() for st in line[ "block" ] ] )
    
        if 'list' in line and bullet_string != "" : # new list point encountered
            text += "\n{}\n".format(bullet_string)
            bullet_string = ""

        if "list" in line and int ( line['list']['level'] ) == int (level+1 ): # entering nested
            line_text = ""
            new_idx,temp_string = list_formatting( line_list[ idx: ] , ( level +1 ) )
            idx = idx + new_idx
            text += temp_string
            
        bullet_string = bullet_string + " " +line_text

        if "list" in line and line['list']['level'] == level-1 : # exiting nested
            return idx , text
        
        idx += 1
    if bullet_string != "" :
        text += "\n{}\n".format(bullet_string)
    return idx , text

def _docx_formating(json_response, index=0, document=None):
    """
    Helping function for docx formating
    """
    if document == None:
        document = docx.Document()
    document.add_heading(f"page_{str(index+1)}",level=6)

    if len(json_response) == 1:
        string_data = json_response[0]["content"]["error"]
        document.add_paragraph(f"{string_data.strip()}")
        return document

    if not isinstance(json_response, list):
        string_data = "Input does not match the expected format"
        document.add_paragraph(f"{string_data}")
        return document

    word_dic_US = enchant.Dict("en_US")
    word_dic_UK = enchant.Dict("en_UK")

    if len(json_response) > 2 :
            two_col = True
    else :
        two_col = False
       
    text_split_flag = False # if word was split in two text lines with "-"
       
    for data in json_response[:2]:
        try:
            if data["idx"] == -1:
                continue
            para_index = 0
            para_string = ""

            if two_col :
                document.add_heading(f'Column {str(data["idx"])}',level=6)

            line_list=[]
            
            for index1,i in enumerate(data["content"]):

                if ( i['alignment'] != None and i['alignment']=='list') :
                        line_list.append(i)
                        continue
                if line_list != []:
                    if len(para_string.strip()) > 2:
                        document.add_paragraph(f"{para_string.strip()}")
                    line_text = list_formatting( line_list )[1]
                    document.add_paragraph(f"{line_text.strip()}")
                    line_list = []
                    para_string = ""

                if i["type"] == "table":
                    if len(para_string.strip()) > 2:
                        document.add_paragraph(f"{para_string.strip()}")
                    table_data = i["block"][0]["content"]
                    df = pd.DataFrame(table_data)
                    t = document.add_table(df.shape[0] + 1, df.shape[1])

                    # add the header rows.
                    for row in range(df.shape[-1]):
                        t.cell(0, row).text = str(df.columns[row])

                    # add the rest of the data frame
                    for row in range(df.shape[0]):
                        for col in range(df.shape[-1]):
                            t.cell(row + 1, col).text = str(df.values[row, col])
                    para_string = ""

                if i["type"] == "p":
                    if i["paragraph_id"] == para_index:
                        string = " ".join([st["content"].strip() for st in i["block"]])
                        if i['alignment'] == 'new_line':
                                string = "\n" + string +"\n"

                        if text_split_flag : # the last ele had a word split
                                para_string = para_string[:-1] + string # removing "-" and concatenating two lines.
                                text_split_flag = False
                        else:
                            para_string = para_string + " " +string # space between lines of a para
                    else:
                        if len(para_string.strip()) > 2:
                            document.add_paragraph(f"{para_string.strip()}")
                        para_index = i["paragraph_id"]
                        string = " ".join([st["content"].strip() for st in i["block"]])
                        para_string = string

                    if string.endswith("-") : # handling words splited by line break ( separated by "-")
                        if index1+1 < len( data["content"] ) and data["content"][index1+1]["type"]== "p" and data["content"][index1+1]["paragraph_id"] == para_index : # next ele belongs to same para
                            next_line =  " ".join([st["content"].strip() for st in data["content"][index1+1]["block"]])
                            try:
                                next_word = next_line.strip()[ : next_line.strip().index(" ")] # word portion that was pushed to next line ( text after '-' ) (first word of next line)
                            except:
                                next_word= next_line # single word
                            try:
                                prev_word = string[ string.rindex(" ")+1: ] # first part of the word along with '-' (last word of current line)
                            except :
                                prev_word = string # single word
                            
                            word_splitted = ( prev_word[:-1] + next_word ).lower()
                            if word_splitted[-1] in string_library.punctuation :
                                word_splitted = word_splitted[:-1].lower()
                            if word_dic_US.check( word_splitted ) or  word_dic_UK.check( word_splitted ): # if after removal of "-" a meaningful word is formed
                                text_split_flag = True

                if i["type"] == "title":
                    if len(para_string.strip()) > 2:
                        document.add_paragraph(f"{para_string.strip()}")
                    string = " ".join([st["content"].strip() for st in i["block"]])
                    if i['font_size'] == "h6" :
                            document.add_heading(f"{string.strip()}",level=6)
                    else :
                        document.add_heading(f"{string.strip()}",level=1)
                    para_string = ""

                if i["type"] == "image":
                    if len(para_string.strip()) > 2:
                        document.add_paragraph(f"{para_string.strip()}")
                    string = i["block"][0]["content"]
                    logo_file = BytesIO(base64.b64decode(string))
                    document.add_picture(
                        logo_file, width=Inches(5.5), height=Inches(int(5.5))
                    )
                    para_string = ""
            paragraph = document.add_paragraph(para_string.strip())
            if line_list != []:
                line_text = list_formatting( line_list )[1]
                document.add_paragraph(f"{line_text.strip()}")
            paragraph.line_spacing_rule = 1
        
        except Exception as ex:
            print("DOCX PARSER: FAILING Page",index+1, ex )
            document.add_heading(f"DOCX RECONSTRUCTION FAILING",level =6)
            document.add_heading(f"Raw Text",level =6)
            document.add_paragraph ( "\n".join( json_response[ -1 ][ "content" ][ "full_text" ] ) )
    return document

def docx_formating(json_response):
    """
    Layout ocr into docx format
    """
    if "math" in json_response :
        html = html_formating(json_response)
        random_name = "temp" + str( time.time() ).replace(".","_") # unique name

        with open ("./{}.html".format(random_name), "w+", encoding='utf8') as f:
            f.write(html)
        f.close()

        cmd = "pandoc -s ./{}.html -o ./{}.docx".format(random_name,random_name)
        os.system( cmd )
        print("pandoc called")
        doc = docx.Document( ('./{}.docx'.format(random_name)) )
        os.remove( ("{}.html".format(random_name)) )
        os.remove( ("{}.docx".format(random_name)) )
    else :
        doc = None
        for index, response in json_response.items():
            doc = _docx_formating(response, int(index), doc)
    return doc